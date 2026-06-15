import os, json, uuid, re, math, io, difflib, html
from datetime import datetime
from pathlib import Path
from typing import Optional, List
from collections import defaultdict

from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse, HTMLResponse
from pydantic import BaseModel
import requests as req_lib

try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data"
PAPERS_DIR = DATA_DIR / "papers"
VECTORS_DIR = DATA_DIR / "vectors"
DB_FILE = DATA_DIR / "db" / "documents.json"
ANALYTICS_FILE = DATA_DIR / "db" / "analytics.json"
AUDIT_FILE = DATA_DIR / "db" / "audit.json"
BOOKMARKS_FILE = DATA_DIR / "db" / "bookmarks.json"
ANNOTATIONS_FILE = DATA_DIR / "db" / "annotations.json"
HEALTH_CACHE_FILE = DATA_DIR / "db" / "health_cache.json"
TIMELINE_FILE = DATA_DIR / "db" / "timeline.json"

for d in [PAPERS_DIR, VECTORS_DIR, DATA_DIR / "db"]:
    d.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="TCG KnowledgeIQ")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

def load_db():
    if DB_FILE.exists():
        try: return json.loads(DB_FILE.read_text())
        except: pass
    return {"documents": []}

def save_db(data):
    DB_FILE.write_text(json.dumps(data, indent=2, default=str))

def load_json_file(path, default):
    if path.exists():
        try: return json.loads(path.read_text())
        except: pass
    return default

def save_json_file(path, data):
    path.write_text(json.dumps(data, indent=2, default=str))

def log_analytics(query_type):
    data = {"total_queries": 0, "summaries": 0, "comparisons": 0}
    if ANALYTICS_FILE.exists():
        try: data = json.loads(ANALYTICS_FILE.read_text())
        except: pass
    data[query_type] = data.get(query_type, 0) + 1
    ANALYTICS_FILE.write_text(json.dumps(data, indent=2))

def get_analytics():
    data = {"total_queries": 0, "summaries": 0, "comparisons": 0}
    if ANALYTICS_FILE.exists():
        try: data = json.loads(ANALYTICS_FILE.read_text())
        except: pass
    return data

def log_audit(action, details):
    audit = []
    if AUDIT_FILE.exists():
        try: audit = json.loads(AUDIT_FILE.read_text())
        except: pass
    audit.insert(0, {"timestamp": datetime.now().isoformat(), "action": action, "details": details})
    # Keep last 1000 logs
    audit = audit[:1000]
    AUDIT_FILE.write_text(json.dumps(audit, indent=2))

def log_timeline(event_type, content, meta=None):
    timeline = load_json_file(TIMELINE_FILE, [])
    timeline.insert(0, {
        "id": str(uuid.uuid4()),
        "timestamp": datetime.now().isoformat(),
        "type": event_type,
        "content": content,
        "meta": meta or {}
    })
    timeline = timeline[:200]
    save_json_file(TIMELINE_FILE, timeline)

class LocalVectorEngine:
    def __init__(self):
        self.index, self.idf, self.doc_count = {}, {}, 0
        self._load_index()

    def _tokenize(self, text):
        text = re.sub(r'[^\w\s]', ' ', text.lower())
        tokens = text.split()
        stops = {'the','a','an','is','it','in','on','at','to','for','of','and'}
        return [t for t in tokens if t not in stops and len(t) > 2]

    def add_document(self, doc_id, text):
        tokens = self._tokenize(text)
        tf = defaultdict(int)
        for t in tokens: tf[t] += 1
        total = max(len(tokens), 1)
        self.index[doc_id] = {"tokens": list(set(tokens)), "tf": {k: v/total for k,v in tf.items()}}
        self.doc_count = len(self.index)
        self._recompute_idf()
        self._save_index()

    def remove_document(self, doc_id):
        if doc_id in self.index:
            del self.index[doc_id]
            self.doc_count = len(self.index)
            self._recompute_idf()
            self._save_index()

    def _recompute_idf(self):
        df = defaultdict(int)
        for doc in self.index.values():
            for t in doc["tokens"]: df[t] += 1
        N = max(self.doc_count, 1)
        self.idf = {t: math.log((N+1)/(df[t]+1))+1 for t in df}

    def search(self, query, top_k=5):
        q_tokens = self._tokenize(query)
        if not q_tokens or not self.index: return []
        scores = {}
        for doc_id, doc in self.index.items():
            score = sum(doc["tf"].get(qt,0) * self.idf.get(qt,0) for qt in q_tokens)
            if score > 0: scores[doc_id] = score
        if not scores: return []
        max_s = max(scores.values())
        res = [{"doc_id": k, "score": round(v/max_s,4)} for k,v in scores.items()]
        res.sort(key=lambda x: x["score"], reverse=True)
        return res[:top_k]

    def get_excerpt(self, doc_id, query):
        vec_file = VECTORS_DIR / f"{doc_id}.txt"
        if not vec_file.exists(): return ""
        text = vec_file.read_text(encoding="utf-8", errors="replace")
        q_tokens = self._tokenize(query)
        words = text.split()
        best_pos, best_score = 0, 0
        for i in range(len(words)):
            chunk = " ".join(words[i:i+50]).lower()
            s = sum(1 for qt in q_tokens if qt in chunk)
            if s > best_score: best_score, best_pos = s, i
        return " ".join(words[max(0,best_pos-10):best_pos+60])

    def _save_index(self):
        (VECTORS_DIR / "_index.json").write_text(json.dumps(self.index, indent=2))

    def _load_index(self):
        idx_file = VECTORS_DIR / "_index.json"
        if idx_file.exists():
            try:
                self.index = json.loads(idx_file.read_text())
                self.doc_count = len(self.index)
                self._recompute_idf()
            except: self.index = {}

vector_engine = LocalVectorEngine()

def extract_text(file_path):
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".pdf" and HAS_PYMUPDF:
            doc = fitz.open(str(file_path))
            return "\n".join(page.get_text() for page in doc)
        elif suffix in (".doc", ".docx") and HAS_DOCX:
            return "\n".join(p.text for p in DocxDocument(str(file_path)).paragraphs)
        else: return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception as e: return f"[Error: {str(e)}]"

def detect_version(filename):
    match = re.match(r'^(.*?)[ _-]v(\d+)(\.[^.]+)?$', filename, re.IGNORECASE)
    if match: return match.group(1).strip(), f"v{match.group(2)}"
    return None, None

# -----------------------------------------------------------------------
# Pydantic Models
# -----------------------------------------------------------------------
class QueryRequest(BaseModel):
    query: str
    top_k: int = 5
    filter_type: Optional[str] = "all"
    language: str = "English"
    compliance_mode: bool = False

class SummarizeRequest(BaseModel):
    doc_id: str
    format: str

class CompareRequest(BaseModel):
    doc1_id: str
    doc2_id: str

class ExportRequest(BaseModel):
    text: str
    filetype: str
    title: str

class ProtocolVerifyRequest(BaseModel):
    query: str
    doc_ids: Optional[List[str]] = None

class ExtractChemicalRequest(BaseModel):
    doc_id: str

class FollowupRequest(BaseModel):
    query: str
    answer: str

class VersionCompareRequest(BaseModel):
    base_name: str
    version1_id: str
    version2_id: str

class BookmarkRequest(BaseModel):
    query: str
    note: Optional[str] = ""

class BookmarkUpdateRequest(BaseModel):
    note: str

class AnnotationRequest(BaseModel):
    doc_id: str
    result_excerpt: str
    note: str

class HealthScoreRequest(BaseModel):
    doc_id: str

class BatchExportRequest(BaseModel):
    doc_ids: List[str]
    title: str
    filetype: str

class CustomAuditRequest(BaseModel):
    action: str
    details: dict

# -----------------------------------------------------------------------
# Core Endpoints
# -----------------------------------------------------------------------

@app.get("/api/documents")
def list_documents():
    return {"documents": load_db()["documents"]}

@app.post("/api/documents/upload")
async def upload_documents(files: List[UploadFile] = File(...)):
    db = load_db()
    for file in files:
        doc_id = str(uuid.uuid4())
        suffix = Path(file.filename).suffix.lower()
        file_path = PAPERS_DIR / f"{doc_id}{suffix}"
        content = await file.read()
        file_path.write_bytes(content)
        text = extract_text(file_path)
        (VECTORS_DIR / f"{doc_id}.txt").write_text(text, encoding="utf-8")
        vector_engine.add_document(doc_id, text)
        base_name, version_label = detect_version(file.filename)
        doc_record = {
            "id": doc_id, "original_name": file.filename, "file_type": suffix.lstrip("."),
            "uploaded_at": datetime.now().isoformat(), "size": len(content),
            "version_base": base_name, "version_label": version_label
        }
        db["documents"].append(doc_record)
        log_timeline("upload", f"Uploaded: {file.filename}", {"doc_id": doc_id, "size": len(content)})
    save_db(db)
    return {"status": "ok", "uploaded": len(files)}

@app.get("/api/documents/{doc_id}/view")
def view_document(doc_id: str, highlight: Optional[str] = Query(None)):
    db = load_db()
    doc = next((d for d in db["documents"] if d["id"] == doc_id), None)
    if not doc: raise HTTPException(404, "Document not found")
    file_path = PAPERS_DIR / f"{doc_id}.{doc['file_type']}"
    if not file_path.exists(): raise HTTPException(404, "File missing")
    if highlight:
        text = extract_text(file_path)
        escaped_text = html.escape(text)
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{doc['original_name']} - KnowledgeIQ View</title>
            <style>
                body {{ background: #0b162c; color: #e2e8f0; font-family: 'Courier New', Courier, monospace; padding: 40px; line-height: 1.8; max-width: 1200px; margin: 0 auto; }}
                h2 {{ font-family: sans-serif; color: #e0b45c; border-bottom: 1px solid rgba(224,180,92,0.3); padding-bottom: 12px; margin-bottom: 24px; }}
                pre {{ white-space: pre-wrap; word-wrap: break-word; font-size: 14.5px; }}
                mark {{ background-color: #e0b45c !important; color: #000000 !important; padding: 2px 6px; border-radius: 4px; font-weight: 900; box-shadow: 0 0 8px rgba(224, 180, 92, 0.6); transition: all 0.3s ease; }}
                .pulse-mark {{ background-color: #fde68a !important; box-shadow: 0 0 16px rgba(224, 180, 92, 1); transform: scale(1.02); }}
            </style>
        </head>
        <body>
            <h2>{doc['original_name']}</h2>
            <pre id="content">{escaped_text}</pre>
            <script>
            (function(){{
                let term = {json.dumps(highlight)};
                if(!term) return;
                let regex = new RegExp(`(${{term.replace(/[.*+?^${{}}()|[\\]\\\\]/g, '\\\\$&')}})`, 'gi');
                let content = document.getElementById('content');
                let html = content.innerHTML;
                let newHtml = html.replace(regex, '<mark>$1</mark>');
                content.innerHTML = newHtml;
                let firstMark = document.querySelector('mark');
                if(firstMark){{
                    firstMark.scrollIntoView({{behavior:'smooth', block:'center'}});
                    firstMark.classList.add('pulse-mark');
                    setTimeout(()=>{{ firstMark.classList.remove('pulse-mark'); }}, 1500);
                }}
            }})();
            </script>
        </body>
        </html>
        """
        return HTMLResponse(content=html_content)
    mime = {"pdf":"application/pdf","txt":"text/plain","html":"text/html"}
    return FileResponse(file_path, media_type=mime.get(doc["file_type"], "application/octet-stream"), headers={"Content-Disposition": "inline"})

@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: str):
    db = load_db()
    db["documents"] = [d for d in db["documents"] if d["id"] != doc_id]
    save_db(db)
    vector_engine.remove_document(doc_id)
    for p in PAPERS_DIR.glob(f"{doc_id}.*"): p.unlink(missing_ok=True)
    (VECTORS_DIR / f"{doc_id}.txt").unlink(missing_ok=True)
    health_cache = load_json_file(HEALTH_CACHE_FILE, {})
    health_cache.pop(doc_id, None)
    save_json_file(HEALTH_CACHE_FILE, health_cache)
    annotations = load_json_file(ANNOTATIONS_FILE, [])
    annotations = [a for a in annotations if a.get("doc_id") != doc_id]
    save_json_file(ANNOTATIONS_FILE, annotations)
    return {"status": "deleted"}

@app.post("/api/query_stream")
async def query_stream(req: QueryRequest):
    log_analytics("total_queries")
    log_timeline("query", req.query, {"filter": req.filter_type, "language": req.language})
    
    # Real-time exhaustive audit logging
    log_audit("search_query", {
        "query": req.query, 
        "filter_type": req.filter_type, 
        "compliance_mode": req.compliance_mode
    })

    db = load_db()
    candidate_docs = db["documents"]
    if req.filter_type and req.filter_type != "all":
        if req.filter_type == "pdf": candidate_docs = [d for d in candidate_docs if d["file_type"] == "pdf"]
        elif req.filter_type == "word": candidate_docs = [d for d in candidate_docs if d["file_type"] in ("doc","docx")]
    candidate_ids = {d["id"] for d in candidate_docs}
    raw_results = [r for r in vector_engine.search(req.query, top_k=req.top_k) if r["doc_id"] in candidate_ids]
    doc_map = {d["id"]: d for d in db["documents"]}
    enriched, context_blocks = [], []
    for r in raw_results:
        doc = doc_map.get(r["doc_id"])
        excerpt = vector_engine.get_excerpt(r["doc_id"], req.query)
        full_text = (VECTORS_DIR / f"{r['doc_id']}.txt").read_text(encoding="utf-8", errors="replace")
        context_blocks.append(f'[Doc: "{doc["original_name"]}"]\n{full_text[:2000]}')
        enriched.append({"doc_id": r["doc_id"], "doc_name": doc["original_name"], "score": r["score"], "excerpt": excerpt, "file_type": doc.get("file_type")})

    async def event_stream():
        yield f"data: {json.dumps({'type': 'results', 'data': enriched})}\n\n"
        if not enriched:
            sys_prompt = f"You are an expert AI. Respond in {req.language}. No documents matched. Use general knowledge. Start with: '⚠️ No internal documents matched. Based on general knowledge: '"
            usr_prompt = req.query
        else:
            sys_prompt = f"You are an expert AI. Respond in {req.language}. Analyze ONLY the provided documents. Write detailed paragraphs, use bullet points, bold text. ALWAYS cite [Doc: 'Name'] inline."
            usr_prompt = f"Question: {req.query}\n\nDocuments:\n" + "\n\n---\n\n".join(context_blocks)
        try:
            r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": f"System: {sys_prompt}\n\nUser: {usr_prompt}", "stream": True}, stream=True, timeout=120)
            for line in r.iter_lines():
                if line:
                    chunk = json.loads(line).get("response", "")
                    yield f"data: {json.dumps({'type': 'chunk', 'text': chunk})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    return StreamingResponse(event_stream(), media_type="text/event-stream")

@app.post("/api/protocol_verify")
def protocol_verify(req: ProtocolVerifyRequest):
    db = load_db()
    forbidden_keywords = ["forbidden_chemical", "unapproved_solvent", "prohibited_step"]
    safety_missing = ["MSDS", "risk assessment", "personal protective equipment"]
    warnings = []
    doc_ids = req.doc_ids or [d["id"] for d in db["documents"]]
    for doc_id in doc_ids:
        text_file = VECTORS_DIR / f"{doc_id}.txt"
        if not text_file.exists(): continue
        text = text_file.read_text(encoding="utf-8", errors="replace").lower()
        doc_name = next((d["original_name"] for d in db["documents"] if d["id"] == doc_id), doc_id)
        for kw in forbidden_keywords:
            if kw in text: warnings.append(f"Document '{doc_name}' contains forbidden term: {kw}")
        for safety in safety_missing:
            if safety not in text: warnings.append(f"Document '{doc_name}' missing safety keyword: {safety}")
    if req.query:
        try:
            llm_check = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": f"System: Compliance officer. User asks: '{req.query}'. List potential compliance violations. Respond with short bullet points only, each under 100 chars. If none, say 'NO_VIOLATIONS'.", "stream": False}, timeout=60)
            llm_response = llm_check.json().get("response", "")
            if "NO_VIOLATIONS" not in llm_response and len(llm_response) < 300: warnings.append(llm_response)
        except: pass
    log_audit("compliance_check", {"query": req.query, "warnings": warnings})
    return {"warnings": warnings}

@app.post("/api/extract_chemical_data")
def extract_chemical_data(req: ExtractChemicalRequest):
    text_file = VECTORS_DIR / f"{req.doc_id}.txt"
    if not text_file.exists(): raise HTTPException(404, "Document not found")
    text = text_file.read_text(encoding="utf-8", errors="replace")[:4000]
    prompt = f"""Extract tabular chemical data (MW, concentration, pH, volume). Return as HTML table. If none, return '<p>No structured chemical data found.</p>'. Text: {text}"""
    try:
        r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False}, timeout=90)
        html_out = r.json().get("response", "<p>Extraction failed.</p>")
        if not html_out.startswith("<table"): html_out = f"<table>{html_out}</td>"
        return {"html_table": html_out}
    except Exception as e: return {"html_table": f"<p>Error: {e}</p>"}

@app.post("/api/entity_graph")
def entity_graph(req: ExtractChemicalRequest):
    text_file = VECTORS_DIR / f"{req.doc_id}.txt"
    if not text_file.exists(): raise HTTPException(404, "Document not found")
    text = text_file.read_text(encoding="utf-8", errors="replace")[:4000]
    prompt = f"""Extract entities (chemicals, targets, parameters) and relationships. Return JSON: {{"nodes":[{{"id":"name","group":"chemical/target/parameter"}}], "links":[{{"source":"from","target":"to","value":"relationship"}}]}}. Text: {text}"""
    try:
        r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False, "format": "json"}, timeout=90)
        return json.loads(r.json().get("response", '{"nodes":[], "links":[]}'))
    except: return {"nodes": [], "links": []}

@app.post("/api/hypothesis")
def generate_hypothesis(req: FollowupRequest):
    prompt = f"""Suggest 2-3 testable scientific hypotheses from the answer. Return JSON list of strings. Query: {req.query}\nAnswer: {req.answer[:2000]}"""
    try:
        r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False, "format": "json"}, timeout=60)
        hyps = json.loads(r.json().get("response", '["Hypothesis 1","Hypothesis 2"]'))
        if not isinstance(hyps, list): hyps = []
        return {"hypotheses": hyps[:3]}
    except: return {"hypotheses": ["Increase flow rate may improve resolution", "Higher temperature could reduce retention time"]}

@app.post("/api/followup")
def generate_followup(req: FollowupRequest):
    prompt = f"""Suggest 3 follow-up questions a researcher would ask. Return JSON list of strings. Query: {req.query}\nAnswer: {req.answer[:2000]}"""
    try:
        r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False, "format": "json"}, timeout=60)
        sugg = json.loads(r.json().get("response", '["What are the risks?","Show related protocols","Extract chemical data"]'))
        if not isinstance(sugg, list): sugg = []
        return {"suggestions": sugg[:3]}
    except: return {"suggestions": ["Check compliance for this procedure", "Extract chemical data", "Compare with earlier version"]}

@app.get("/api/versions")
def get_version_groups():
    db = load_db()
    groups = defaultdict(list)
    for doc in db["documents"]:
        base = doc.get("version_base")
        if base: groups[base].append({"id": doc["id"], "original_name": doc["original_name"], "version_label": doc.get("version_label", "unknown")})
    for base in groups: groups[base].sort(key=lambda x: x["version_label"])
    return {"groups": [{"base_name": base, "versions": vers} for base, vers in groups.items()]}

@app.post("/api/compare_versions")
def compare_versions(req: VersionCompareRequest):
    t1 = VECTORS_DIR / f"{req.version1_id}.txt"
    t2 = VECTORS_DIR / f"{req.version2_id}.txt"
    if not t1.exists() or not t2.exists(): raise HTTPException(404, "Version texts missing")
    lines1 = t1.read_text(encoding="utf-8", errors="replace").splitlines()
    lines2 = t2.read_text(encoding="utf-8", errors="replace").splitlines()
    diff = difflib.HtmlDiff(wrapcolumn=80)
    html_diff = diff.make_table(lines1, lines2, fromdesc=req.version1_id, todesc=req.version2_id)
    log_audit("version_compare", {"version1": req.version1_id, "version2": req.version2_id})
    return {"diff_html": html_diff}

@app.get("/api/audit")
def get_audit():
    if AUDIT_FILE.exists(): return {"audit": json.loads(AUDIT_FILE.read_text())}
    return {"audit": []}

@app.post("/api/audit")
def post_audit(req: CustomAuditRequest):
    log_audit(req.action, req.details)
    return {"status": "ok"}

@app.delete("/api/audit")
def clear_audit():
    save_json_file(AUDIT_FILE, [])
    return {"status": "cleared"}

@app.post("/api/summarize")
def summarize_doc(req: SummarizeRequest):
    log_analytics("summaries")
    db = load_db()
    doc_name = next((d["original_name"] for d in db["documents"] if d["id"] == req.doc_id), req.doc_id)
    log_timeline("summary", f"Summarized: {doc_name}", {"doc_id": req.doc_id, "format": req.format})
    text_file = VECTORS_DIR / f"{req.doc_id}.txt"
    if not text_file.exists(): return {"error": "Document not found."}
    text = text_file.read_text(encoding="utf-8", errors="replace")[:6000]
    prompt = f"System: Expert analyst. Provide detailed {req.format} format response.\n\nDocument:\n{text}"
    try:
        r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False}, timeout=120)
        return {"summary": r.json().get("response", "")}
    except: return {"summary": "Error: Ensure Ollama is running."}

@app.post("/api/compare")
def compare_docs(req: CompareRequest):
    log_analytics("comparisons")
    db = load_db()
    doc1_name = next((d["original_name"] for d in db["documents"] if d["id"] == req.doc1_id), req.doc1_id)
    doc2_name = next((d["original_name"] for d in db["documents"] if d["id"] == req.doc2_id), req.doc2_id)
    log_timeline("compare", f"Compared: {doc1_name} vs {doc2_name}", {"doc1": req.doc1_id, "doc2": req.doc2_id})
    t1 = VECTORS_DIR / f"{req.doc1_id}.txt"
    t2 = VECTORS_DIR / f"{req.doc2_id}.txt"
    if not t1.exists() or not t2.exists(): return {"error": "Documents not found."}
    text1 = t1.read_text(encoding="utf-8", errors="replace")[:3000]
    text2 = t2.read_text(encoding="utf-8", errors="replace")[:3000]
    prompt = f"System: Expert analyst. Compare documents deeply, highlight similarities, differences, conflicts. Use bullet points and bold.\n\nDoc1:\n{text1}\n\nDoc2:\n{text2}"
    try:
        r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False}, timeout=120)
        return {"comparison": r.json().get("response", "")}
    except: return {"comparison": "Error: Ensure Ollama is running."}

@app.post("/api/export")
def export_document(req: ExportRequest):
    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', req.text)
    clean = re.sub(r'#+\s', '', clean)
    if req.filetype == "docx" and HAS_DOCX:
        doc = DocxDocument()
        doc.add_heading(req.title, 0)
        doc.add_paragraph(clean)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={req.title.replace(' ', '_')}.docx"})
    elif req.filetype == "pdf" and HAS_FPDF:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", style="B", size=16)
        pdf.cell(0, 10, text=req.title.encode('latin-1', 'replace').decode('latin-1'), align='C')
        pdf.ln(10)
        pdf.set_font("helvetica", size=12)
        pdf.multi_cell(0, 7, text=clean.encode('latin-1', 'replace').decode('latin-1'))
        buf = io.BytesIO(pdf.output())
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={req.title.replace(' ', '_')}.pdf"})
    else:
        buf = io.BytesIO(clean.encode('utf-8'))
        buf.seek(0)
        return StreamingResponse(buf, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename={req.title.replace(' ', '_')}.txt"})

@app.get("/api/analytics")
def analytics_data():
    db = load_db()
    stats = get_analytics()
    total_size = sum(d.get("size", 0) for d in db["documents"])
    return {
        "total_documents": len(db["documents"]),
        "total_queries": stats.get("total_queries", 0),
        "summaries_run": stats.get("summaries", 0),
        "comparisons_run": stats.get("comparisons", 0),
        "vault_size_mb": round(total_size / (1024 * 1024), 2)
    }

@app.get("/api/bookmarks")
def list_bookmarks():
    return {"bookmarks": load_json_file(BOOKMARKS_FILE, [])}

@app.post("/api/bookmarks")
def add_bookmark(req: BookmarkRequest):
    bookmarks = load_json_file(BOOKMARKS_FILE, [])
    if any(b["query"] == req.query for b in bookmarks):
        return {"status": "already_exists"}
    bookmark = {
        "id": str(uuid.uuid4()),
        "query": req.query,
        "note": req.note or "",
        "created_at": datetime.now().isoformat()
    }
    bookmarks.insert(0, bookmark)
    save_json_file(BOOKMARKS_FILE, bookmarks)
    log_audit("bookmark_added", {"query": req.query})
    return {"status": "ok", "bookmark": bookmark}

@app.patch("/api/bookmarks/{bookmark_id}")
def update_bookmark_note(bookmark_id: str, req: BookmarkUpdateRequest):
    bookmarks = load_json_file(BOOKMARKS_FILE, [])
    for b in bookmarks:
        if b["id"] == bookmark_id:
            b["note"] = req.note
            save_json_file(BOOKMARKS_FILE, bookmarks)
            return {"status": "ok"}
    raise HTTPException(404, "Bookmark not found")

@app.delete("/api/bookmarks/{bookmark_id}")
def delete_bookmark(bookmark_id: str):
    bookmarks = load_json_file(BOOKMARKS_FILE, [])
    bookmarks = [b for b in bookmarks if b["id"] != bookmark_id]
    save_json_file(BOOKMARKS_FILE, bookmarks)
    return {"status": "deleted"}

@app.post("/api/health_score")
def compute_health_score(req: HealthScoreRequest):
    health_cache = load_json_file(HEALTH_CACHE_FILE, {})
    if req.doc_id in health_cache:
        return health_cache[req.doc_id]

    text_file = VECTORS_DIR / f"{req.doc_id}.txt"
    if not text_file.exists(): raise HTTPException(404, "Document not found")
    text = text_file.read_text(encoding="utf-8", errors="replace")
    words = text.split()
    word_count = len(words)

    issues = []
    scores = {}

    completeness = min(25, int(word_count / 40))
    section_keywords = ["introduction", "method", "result", "conclusion", "abstract", "objective", "summary", "procedure"]
    found_sections = sum(1 for kw in section_keywords if kw in text.lower())
    completeness = min(25, completeness + found_sections * 2)
    scores["completeness"] = completeness
    if word_count < 200: issues.append("Document is very short (< 200 words)")
    if found_sections < 2: issues.append("Missing key structural sections (intro, method, results)")

    compliance_terms = ["msds", "risk assessment", "ppe", "hazard", "sop", "gmp", "fda", "ich", "coa", "specification"]
    found_compliance = sum(1 for t in compliance_terms if t in text.lower())
    compliance_score = min(25, found_compliance * 3)
    scores["compliance"] = compliance_score
    if found_compliance < 3: issues.append(f"Low compliance terminology ({found_compliance}/10 keywords found)")

    sentences = re.split(r'[.!?]+', text)
    sentences = [s for s in sentences if len(s.strip()) > 10]
    avg_sent_len = sum(len(s.split()) for s in sentences) / max(len(sentences), 1)
    readability = 25 if avg_sent_len < 20 else max(0, 25 - int((avg_sent_len - 20) * 1.5))
    scores["readability"] = readability
    if avg_sent_len > 35: issues.append(f"Average sentence length is high ({int(avg_sent_len)} words) — hard to read")

    number_pattern = re.compile(r'\b\d+[\.,]?\d*\s*(%|mg|mL|mM|µM|g/L|pH|nm|°C|min|hr|h)\b')
    data_hits = len(number_pattern.findall(text))
    data_richness = min(25, data_hits * 2)
    scores["data_richness"] = data_richness
    if data_hits < 5: issues.append("Very few quantitative data points (numbers with units) detected")

    total = completeness + compliance_score + readability + data_richness
    grade = "high" if total >= 70 else "mid" if total >= 40 else "low"

    llm_issues = []
    try:
        snippet = text[:1500]
        prompt = f"""You are a document quality analyst. Read this document excerpt and list up to 3 specific quality issues (missing data, unclear instructions, compliance gaps). Each issue must be a single sentence under 80 chars. Return as JSON list of strings only.\n\nDocument:\n{snippet}"""
        r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False, "format": "json"}, timeout=45)
        llm_issues = json.loads(r.json().get("response", "[]"))
        if not isinstance(llm_issues, list): llm_issues = []
        llm_issues = [i for i in llm_issues if isinstance(i, str) and len(i) < 150][:3]
    except: pass

    result = {
        "doc_id": req.doc_id,
        "total_score": total,
        "grade": grade,
        "breakdown": scores,
        "issues": issues + llm_issues,
        "word_count": word_count,
        "computed_at": datetime.now().isoformat()
    }

    health_cache[req.doc_id] = result
    save_json_file(HEALTH_CACHE_FILE, health_cache)
    log_audit("health_score", {"doc_id": req.doc_id, "score": total, "grade": grade})
    return result

@app.delete("/api/health_score/{doc_id}")
def invalidate_health_cache(doc_id: str):
    health_cache = load_json_file(HEALTH_CACHE_FILE, {})
    health_cache.pop(doc_id, None)
    save_json_file(HEALTH_CACHE_FILE, health_cache)
    return {"status": "cache_cleared"}

@app.get("/api/annotations")
def list_annotations(doc_id: Optional[str] = Query(None)):
    annotations = load_json_file(ANNOTATIONS_FILE, [])
    if doc_id:
        annotations = [a for a in annotations if a.get("doc_id") == doc_id]
    return {"annotations": annotations}

@app.post("/api/annotations")
def add_annotation(req: AnnotationRequest):
    annotations = load_json_file(ANNOTATIONS_FILE, [])
    annotation = {
        "id": str(uuid.uuid4()),
        "doc_id": req.doc_id,
        "result_excerpt": req.result_excerpt[:200],
        "note": req.note,
        "created_at": datetime.now().isoformat()
    }
    annotations.insert(0, annotation)
    save_json_file(ANNOTATIONS_FILE, annotations)
    log_audit("annotation_added", {"doc_id": req.doc_id, "note_preview": req.note[:80]})
    return {"status": "ok", "annotation": annotation}

@app.delete("/api/annotations/{annotation_id}")
def delete_annotation(annotation_id: str):
    annotations = load_json_file(ANNOTATIONS_FILE, [])
    annotations = [a for a in annotations if a["id"] != annotation_id]
    save_json_file(ANNOTATIONS_FILE, annotations)
    return {"status": "deleted"}

@app.get("/api/timeline")
def get_timeline(limit: int = Query(50, le=200)):
    timeline = load_json_file(TIMELINE_FILE, [])
    return {"timeline": timeline[:limit]}

@app.delete("/api/timeline")
def clear_timeline():
    save_json_file(TIMELINE_FILE, [])
    return {"status": "cleared"}

@app.post("/api/batch_export")
def batch_export(req: BatchExportRequest):
    db = load_db()
    doc_map = {d["id"]: d for d in db["documents"]}
    sections = []

    for doc_id in req.doc_ids:
        doc = doc_map.get(doc_id)
        if not doc:
            sections.append({"name": doc_id, "text": "[Document not found]"})
            continue
        text_file = VECTORS_DIR / f"{doc_id}.txt"
        if not text_file.exists():
            sections.append({"name": doc["original_name"], "text": "[Text not available]"})
            continue
        text = text_file.read_text(encoding="utf-8", errors="replace")[:3000]
        summary_words = text.split()[:300]
        summary = " ".join(summary_words)
        try:
            prompt = f"Summarize this document in 2-3 sentences for an executive report. Return plain text only.\n\nDocument ({doc['original_name']}):\n{text[:2000]}"
            r = req_lib.post("http://localhost:11434/api/generate", json={"model": "llama3", "prompt": prompt, "stream": False}, timeout=60)
            llm_summary = r.json().get("response", "").strip()
            if llm_summary: summary = llm_summary
        except: pass
        sections.append({"name": doc["original_name"], "text": summary, "size_kb": round(doc.get("size", 0) / 1024, 1), "uploaded_at": doc.get("uploaded_at", "")})

    log_timeline("batch_export", f"Batch export: {len(sections)} docs", {"format": req.filetype, "doc_count": len(sections)})
    log_audit("batch_export", {"doc_ids": req.doc_ids, "format": req.filetype})

    if req.filetype == "docx" and HAS_DOCX:
        doc_out = DocxDocument()
        doc_out.add_heading(req.title, 0)
        doc_out.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Documents: {len(sections)}")
        doc_out.add_paragraph("")
        for sec in sections:
            doc_out.add_heading(sec["name"], level=1)
            if sec.get("uploaded_at"):
                doc_out.add_paragraph(f"Uploaded: {sec['uploaded_at'][:10]}  |  Size: {sec.get('size_kb', '?')} KB", style="Intense Quote" if "Intense Quote" in [s.name for s in doc_out.styles] else "Normal")
            doc_out.add_paragraph(re.sub(r'\*\*(.*?)\*\*', r'\1', sec["text"]))
            doc_out.add_paragraph("")
        buf = io.BytesIO()
        doc_out.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                  headers={"Content-Disposition": f"attachment; filename={req.title.replace(' ', '_')}.docx"})
    elif req.filetype == "pdf" and HAS_FPDF:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("helvetica", style="B", size=18)
        safe_title = req.title.encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 12, text=safe_title, align='C')
        pdf.ln(6)
        pdf.set_font("helvetica", size=10)
        pdf.cell(0, 8, text=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  {len(sections)} documents", align='C')
        pdf.ln(12)
        for sec in sections:
            pdf.set_font("helvetica", style="B", size=13)
            sec_name = sec["name"].encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 8, text=sec_name)
            pdf.ln(5)
            if sec.get("uploaded_at"):
                pdf.set_font("helvetica", style="I", size=9)
                pdf.cell(0, 6, text=f"Uploaded: {sec['uploaded_at'][:10]}  |  Size: {sec.get('size_kb','?')} KB")
                pdf.ln(4)
            pdf.set_font("helvetica", size=11)
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', sec["text"])
            clean_text = clean_text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 7, text=clean_text)
            pdf.ln(8)
        buf = io.BytesIO(pdf.output())
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/pdf",
                                  headers={"Content-Disposition": f"attachment; filename={req.title.replace(' ', '_')}.pdf"})
    else:
        lines = [req.title, "=" * len(req.title), f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]
        for sec in sections:
            lines.append(f"--- {sec['name']} ---")
            if sec.get("uploaded_at"): lines.append(f"Uploaded: {sec['uploaded_at'][:10]}")
            lines.append(sec["text"])
            lines.append("")
        buf = io.BytesIO("\n".join(lines).encode("utf-8"))
        buf.seek(0)
        return StreamingResponse(buf, media_type="text/plain",
                                  headers={"Content-Disposition": f"attachment; filename={req.title.replace(' ', '_')}.txt"})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)